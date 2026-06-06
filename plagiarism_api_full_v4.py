"""
Plagiarism Detection API v4 — Full Multi-Format + Image Plagiarism
====================================================================
Supporte :
  - Texte brut (.txt, .md)
  - Code source (.py, .js, .java, .c, .cpp...)
  - PDF (.pdf) — extraction texte + images
  - Word (.docx) — extraction texte + images
  - Images directes (.png, .jpg, .jpeg, .gif, .bmp, .webp)
  - ZIP contenant plusieurs fichiers

Detection :
  - Texte : TF-IDF + BERT + Winnowing + LCS + Stop-words FR
  - Code  : TF-IDF + AST + Winnowing + LCS + Normalisation identifiants
  - Images: pHash (copie exacte/modifiée) + Feature comparison (similaire)

Tout retourne du JSON pour Laravel.

Dependances :
  pip install fastapi uvicorn scikit-learn numpy nltk sentence-transformers
  pip install Pillow imagehash python-multipart
  pip install PyMuPDF        (pour PDF — extraction texte + images)
  pip install python-docx     (pour Word — extraction texte + images)
"""

import os
import re
import json
import ast
import hashlib
import difflib
import io
import math
import zipfile
import tempfile
import shutil
from typing import List, Dict, Any, Optional, Tuple
from collections import Counter

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import uvicorn

import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# ===== PIL / ImageHash =====
try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    print("[WARN] Pillow non installe. Detection d'images desactivee.")
    print("       pip install Pillow imagehash")

try:
    import imagehash
    HAS_IMAGEHASH = True
except ImportError:
    HAS_IMAGEHASH = False
    print("[WARN] imagehash non installe. pip install imagehash")

# ===== PyMuPDF (PDF) =====
try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False
    print("[WARN] PyMuPDF non installe. Extraction PDF desactivee.")
    print("       pip install PyMuPDF")

# ===== python-docx (Word) =====
try:
    from docx import Document as DocxDocument
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("[WARN] python-docx non installe. Extraction Word desactivee.")
    print("       pip install python-docx")

# ===========================================================================
# 1. CONFIGURATION
# ===========================================================================
DATA_FILE = "./submissions.json"
UPLOAD_DIR = "uploads"
IMAGES_DIR = os.path.join(UPLOAD_DIR, "images")
MAX_RESULTS = 10
WINNOWING_WINDOW = 4
PARAGRAPH_MIN_LENGTH = 20

# Seuils images
IMAGE_HASH_THRESHOLD = 10       # distance pHash max pour "similaire"
IMAGE_FEATURE_THRESHOLD = 0.80  # similarité cosine sur features

WEIGHTS_TEXT = {"tfidf": 0.15, "semantic": 0.40, "winnowing": 0.20, "lcs": 0.25}
WEIGHTS_CODE = {"tfidf": 0.20, "winnowing": 0.20, "ast": 0.35, "lcs": 0.15, "semantic": 0.10}

LARAVEL_URL = os.environ.get("LARAVEL_URL", "http://localhost:8000")

app = FastAPI(title="Plagiarism Detection API v4", version="4.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Créer les dossiers
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(IMAGES_DIR, exist_ok=True)


# ===========================================================================
# 2. NLP
# ===========================================================================
_nlp_ready = False
_stop_words = set()
_stemmer = None

def init_nlp():
    global _nlp_ready, _stop_words, _stemmer
    if _nlp_ready:
        return
    try:
        import nltk
        try:
            nltk.data.find('corpora/stopwords')
        except LookupError:
            nltk.download('stopwords', quiet=True)
        from nltk.corpus import stopwords
        from nltk.stem import SnowballStemmer
        _stop_words = set(stopwords.words('french')) | set(stopwords.words('english'))
        _stop_words |= set(['le','la','les','un','une','des','du','de','et','est','en',
            'que','qui','dans','ce','il','ne','sur','se','pas','plus','par','je','avec',
            'tout','au','son','cette','mais','sont','aussi','ou','leur','y','a','ete',
            'pour','elle','nous','vous','ils','on','ses','sa','mon','ton','ma','ta',
            'nos','vos','ces','mes','tes','the','is','are','was','were','be','been',
            'have','has','had','do','does','did','will','would','could','should',
            'may','might','must','shall','can','this','that','these','those','it','its'])
        _stemmer = SnowballStemmer('french')
        _nlp_ready = True
        print("[INFO] NLP OK")
    except ImportError:
        print("[WARN] NLTK non installe. pip install nltk")


def stem_text(text):
    return " ".join(_stemmer.stem(w) for w in text.split()) if _stemmer else text

def remove_stop_words(text):
    return " ".join(w for w in text.split() if w.lower() not in _stop_words and len(w)>1) if _stop_words else text


# ===========================================================================
# 3. SENTENCE-TRANSFORMERS
# ===========================================================================
_semantic_model = None

def get_semantic_model():
    global _semantic_model
    if _semantic_model is None:
        try:
            from sentence_transformers import SentenceTransformer
            _semantic_model = SentenceTransformer("all-MiniLM-L6-v2")
            print("[INFO] sentence-transformers OK")
        except ImportError:
            print("[WARN] sentence-transformers non installe.")
            return None
        except Exception as e:
            print(f"[ERROR] modele : {e}")
            return None
    return _semantic_model


# ===========================================================================
# 4. EXTRACTEURS DE FICHIERS — PDF, Word, Images
# ===========================================================================

def extract_pdf(filepath: str) -> Dict:
    """
    Extrait texte et images d'un PDF.
    Retourne : {"text": str, "images": [PIL.Image, ...], "pages": int}
    """
    if not HAS_FITZ:
        return {"text": "", "images": [], "pages": 0, "error": "PyMuPDF non installe"}

    doc = fitz.open(filepath)
    all_text = []
    all_images = []
    num_pages = len(doc)  # sauvegarder AVANT close()

    for page_num in range(num_pages):
        page = doc[page_num]

        # Extraire texte
        text = page.get_text()
        if text.strip():
            all_text.append(f"--- Page {page_num+1} ---\n{text}")

        # Extraire images
        image_list = page.get_images(full=True)
        for img_idx, img in enumerate(image_list):
            try:
                xref = img[0]
                base_image = doc.extract_image(xref)
                if base_image:
                    img_bytes = base_image["image"]
                    img_ext = base_image.get("ext", "png")
                    pil_img = Image.open(io.BytesIO(img_bytes))
                    if pil_img.size[0] > 32 and pil_img.size[1] > 32:  # ignorer trop petit
                        all_images.append(pil_img)
            except Exception:
                continue

    doc.close()
    return {
        "text": "\n\n".join(all_text),
        "images": all_images,
        "pages": num_pages,
    }


def extract_docx(filepath: str) -> Dict:
    """
    Extrait texte et images d'un fichier Word (.docx).
    """
    if not HAS_DOCX:
        return {"text": "", "images": [], "error": "python-docx non installe"}

    doc = DocxDocument(filepath)
    all_text = []

    # Extraire texte de tous les paragraphes
    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(para.text)

    # Extraire texte des tableaux
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text)
            if row_text:
                all_text.append(" | ".join(row_text))

    # Extraire images
    all_images = []
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    img_data = rel.target_part.blob
                    pil_img = Image.open(io.BytesIO(img_data))
                    if pil_img.size[0] > 32 and pil_img.size[1] > 32:
                        all_images.append(pil_img)
                except Exception:
                    continue
    except Exception:
        pass

    return {
        "text": "\n\n".join(all_text),
        "images": all_images,
    }


def extract_file(file_bytes: bytes, filename: str) -> Dict:
    """
    Fonction principale d'extraction. Detecte le type de fichier
    et extrait texte + images.
    Retourne : {"text": str, "images": [PIL.Image], "file_type": str, "metadata": dict}
    """
    ext = os.path.splitext(filename)[1].lower()

    # === PDF ===
    if ext in ('.pdf',):
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            result = extract_pdf(tmp_path)
            file_type = "text"  # on traite le contenu texte du PDF comme du texte
            return {
                "text": result["text"],
                "images": result.get("images", []),
                "file_type": file_type,
                "metadata": {
                    "format": "pdf",
                    "pages": result.get("pages", 0),
                    "images_extracted": len(result.get("images", [])),
                }
            }
        except Exception as e:
            return {"text": "", "images": [], "file_type": "text", "metadata": {"error": str(e)}}
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    # === Word ===
    if ext in ('.docx',):
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            result = extract_docx(tmp_path)
            return {
                "text": result["text"],
                "images": result.get("images", []),
                "file_type": "text",
                "metadata": {
                    "format": "docx",
                    "images_extracted": len(result.get("images", [])),
                }
            }
        except Exception as e:
            return {"text": "", "images": [], "file_type": "text", "metadata": {"error": str(e)}}
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    # === Image directe ===
    if ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.tiff', '.tif'):
        if HAS_PIL:
            try:
                pil_img = Image.open(io.BytesIO(file_bytes))
                # Convertir en RGB si nécessaire
                if pil_img.mode != 'RGB':
                    pil_img = pil_img.convert('RGB')
                return {
                    "text": "",
                    "images": [pil_img],
                    "file_type": "image",
                    "metadata": {
                        "format": ext.lstrip('.'),
                        "width": pil_img.size[0],
                        "height": pil_img.size[1],
                    }
                }
            except Exception as e:
                return {"text": "", "images": [], "file_type": "image", "metadata": {"error": str(e)}}
        return {"text": "", "images": [], "file_type": "image", "metadata": {"error": "Pillow non installe"}}

    # === ZIP ===
    if ext in ('.zip',):
        return extract_zip(file_bytes)

    # === Code / Texte brut ===
    CODE_EXTS = {'.py', '.js', '.jsx', '.ts', '.tsx', '.java', '.c', '.cpp', '.h', '.hpp',
                 '.cs', '.php', '.rb', '.go', '.rs', '.swift', '.kt', '.scala', '.r',
                 '.sql', '.sh', '.bash', '.html', '.css', '.scss', '.xml', '.yaml', '.yml',
                 '.json', '.toml', '.ini', '.cfg', '.conf'}

    try:
        text = file_bytes.decode('utf-8', errors='ignore')
    except Exception:
        text = file_bytes.decode('latin-1', errors='ignore')

    file_type = "code" if ext in CODE_EXTS else "text"
    return {
        "text": text,
        "images": [],
        "file_type": file_type,
        "metadata": {"format": ext.lstrip('.')}
    }


def extract_zip(zip_bytes: bytes) -> Dict:
    """Extrait tous les fichiers d'un ZIP."""
    all_text = []
    all_images = []
    code_exts = {'.py','.js','.jsx','.ts','.tsx','.java','.c','.cpp','.h','.hpp',
                 '.cs','.php','.rb','.go','.rs','.swift','.kt','.html','.css','.sql','.sh'}
    metadata = {"format": "zip", "files": []}

    try:
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
            for name in zf.namelist():
                if name.startswith('__MACOSX') or name.endswith('/'):
                    continue
                try:
                    data = zf.read(name)
                    ext = os.path.splitext(name)[1].lower()

                    if ext in ('.png','.jpg','.jpeg','.gif','.bmp','.webp') and HAS_PIL:
                        try:
                            img = Image.open(io.BytesIO(data))
                            if img.mode != 'RGB':
                                img = img.convert('RGB')
                            all_images.append(img)
                        except Exception:
                            pass
                    elif ext in ('.pdf',) and HAS_FITZ:
                        result = extract_pdf_from_bytes(data)
                        if result["text"]:
                            all_text.append(f"=== {name} ===\n{result['text']}")
                        all_images.extend(result.get("images", []))
                    elif ext in ('.docx',) and HAS_DOCX:
                        result = extract_docx_from_bytes(data)
                        if result["text"]:
                            all_text.append(f"=== {name} ===\n{result['text']}")
                        all_images.extend(result.get("images", []))
                    else:
                        try:
                            txt = data.decode('utf-8', errors='ignore')
                            all_text.append(f"=== {name} ===\n{txt}")
                        except Exception:
                            pass

                    metadata["files"].append(name)
                except Exception:
                    continue
    except Exception as e:
        metadata["error"] = str(e)

    return {
        "text": "\n\n".join(all_text),
        "images": all_images,
        "file_type": "text",
        "metadata": metadata
    }


def extract_pdf_from_bytes(data: bytes) -> Dict:
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        return extract_pdf(tmp_path)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


def extract_docx_from_bytes(data: bytes) -> Dict:
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        return extract_docx(tmp_path)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ===========================================================================
# 5. IMAGE COMPARISON — pHash + Feature-based
# ===========================================================================

def compute_image_hash(img: Image.Image, hash_size: int = 16) -> Optional[str]:
    """Calcule le perceptual hash (pHash) d'une image."""
    if not HAS_IMAGEHASH:
        return None
    try:
        if img.mode != 'RGB':
            img = img.convert('RGB')
        # Redimensionner pour performance (pHash le fait déjà, mais on normalise)
        h = imagehash.phash(img, hash_size=hash_size)
        return str(h)
    except Exception:
        return None


def compute_image_features(img: Image.Image, size: tuple = (64, 64)) -> Optional[np.ndarray]:
    """
    Extrait un vecteur de features basiques d'une image :
    - Histogramme RGB (768 valeurs : 256 par canal)
    - Redimensionné et aplati
    """
    if not HAS_PIL:
        return None
    try:
        if img.mode != 'RGB':
            img = img.convert('RGB')
        img_small = img.resize(size, Image.LANCZOS)
        arr = np.array(img_small)
        # Histogramme par canal
        features = []
        for channel in range(3):
            hist, _ = np.histogram(arr[:,:,channel], bins=64, range=(0, 256))
            features.extend(hist.tolist())
        # Normaliser
        vec = np.array(features, dtype=np.float32)
        norm = np.linalg.norm(vec)
        if norm > 0:
            vec = vec / norm
        return vec
    except Exception:
        return None


def compare_images(new_images: List[Image.Image],
                   existing_image_data: List[Dict]) -> List[Dict]:
    """
    Compare les nouvelles images avec toutes les images de la base.
    Retourne la liste des matches d'images.

    Pour chaque nouvelle image, on compare avec chaque image existante :
    - pHash : distance hamming (seuil = IMAGE_HASH_THRESHOLD)
    - Features : cosine similarity (seuil = IMAGE_FEATURE_THRESHOLD)

    existing_image_data format :
    [{"hash": "abcd...", "features": [0.1, 0.2, ...], "filename": "file.pdf", "image_index": 0}, ...]
    """
    if not new_images or not existing_image_data or not HAS_PIL:
        return []

    matches = []

    for new_idx, new_img in enumerate(new_images):
        new_hash = compute_image_hash(new_img)
        new_features = compute_image_features(new_img)

        for existing in existing_image_data:
            ex_hash = existing.get("hash")
            ex_features_bytes = existing.get("features")

            # Restaurer les features depuis la liste
            if ex_features_bytes:
                ex_features = np.array(ex_features_bytes, dtype=np.float32)
            else:
                ex_features = None

            phash_dist = None
            feature_sim = None

            # Comparaison pHash
            if new_hash and ex_hash and HAS_IMAGEHASH:
                try:
                    h1 = imagehash.hex_to_hash(new_hash, hash_size=16)
                    h2 = imagehash.hex_to_hash(ex_hash, hash_size=16)
                    phash_dist = h1 - h2  # distance de Hamming
                except Exception:
                    try:
                        h1 = imagehash.hex_to_hash(new_hash)
                        h2 = imagehash.hex_to_hash(ex_hash)
                        phash_dist = h1 - h2
                    except Exception:
                        pass

            # Comparaison features
            if new_features is not None and ex_features is not None:
                try:
                    n1 = np.linalg.norm(new_features)
                    n2 = np.linalg.norm(ex_features)
                    if n1 > 0 and n2 > 0:
                        feature_sim = float(np.dot(new_features, ex_features) / (n1 * n2))
                    else:
                        feature_sim = 0.0
                except Exception:
                    pass

            # Déterminer si c'est un match
            is_match = False
            level = "none"
            confidence = 0.0

            if phash_dist is not None and phash_dist <= 5:
                is_match = True
                level = "critical"
                confidence = max(confidence, 1.0 - (phash_dist / 64.0))
            elif phash_dist is not None and phash_dist <= IMAGE_HASH_THRESHOLD:
                is_match = True
                level = "high"
                confidence = max(confidence, 1.0 - (phash_dist / 64.0))
            elif feature_sim is not None and feature_sim >= 0.95:
                is_match = True
                level = "critical"
                confidence = max(confidence, feature_sim)
            elif feature_sim is not None and feature_sim >= IMAGE_FEATURE_THRESHOLD:
                is_match = True
                level = "high" if feature_sim >= 0.90 else "medium"
                confidence = max(confidence, feature_sim)

            if is_match:
                matches.append({
                    "new_image_index": new_idx,
                    "matched_filename": existing.get("filename", "?"),
                    "matched_image_index": existing.get("image_index", 0),
                    "level": level,
                    "confidence": round(confidence, 4),
                    "phash_distance": phash_dist,
                    "feature_similarity": round(feature_sim, 4) if feature_sim else None,
                })

    # Dédoublonner : garder le meilleur match par image
    best_per_image = {}
    for m in matches:
        key = m["new_image_index"]
        if key not in best_per_image or m["confidence"] > best_per_image[key]["confidence"]:
            best_per_image[key] = m

    return sorted(best_per_image.values(), key=lambda x: x["confidence"], reverse=True)


# ===========================================================================
# 6. PRETRAITEMENT TEXTE / CODE
# ===========================================================================

def normalize_identifiers(code: str) -> str:
    id_pattern = re.compile(r'\b[a-zA-Z_][a-zA-Z0-9_]*\b')
    reserved = set([
        'if','else','elif','for','while','def','class','return','import','from','as',
        'try','except','finally','with','yield','lambda','and','or','not','in','is',
        'pass','break','continue','raise','assert','del','global','nonlocal','async','await',
        'int','str','float','bool','list','dict','set','tuple','None','True','False',
        'self','super','print','range','len','type','var','let','const','function',
        'new','this','null','undefined','switch','case','default','typeof','instanceof',
        'public','private','protected','static','final','abstract','interface','extends',
        'implements','package','throws','throw','catch','boolean','char','String','System',
        'include','define','ifdef','endif','struct','enum','typedef','sizeof','printf',
        'scanf','cout','cin','endl','main','args','argv','init','append','extend',
        'insert','remove','pop','sort','map','filter','reduce','input','open',
        'close','read','write',
    ])
    id_map = {}
    counter = [0]
    for token in id_pattern.findall(code):
        if token not in reserved and token not in id_map:
            id_map[token] = f"ID_{counter[0]}"; counter[0] += 1
    return id_pattern.sub(lambda m: id_map.get(m.group(0), m.group(0)), code)


def preprocess_code(code):
    code = re.sub(r'/\*.*?\*/', '', code, flags=re.DOTALL)
    code = re.sub(r'//.*?$', '', code, flags=re.MULTILINE)
    code = re.sub(r'#.*?$', '', code, flags=re.MULTILINE)
    code = re.sub(r'"""[\s\S]*?"""', ' """STR""" ', code)
    code = re.sub(r"'''[\s\S]*?'''", " '''STR''' ", code)
    code = re.sub(r'"[^"\\]*(\\.[^"\\]*)*"', ' "STR" ', code)
    code = re.sub(r"'[^'\\]*(\\.[^'\\]*)*'", " 'STR' ", code)
    code = re.sub(r'\b\d+\.?\d*\b', 'NUM', code)
    code = normalize_identifiers(code)
    return re.sub(r'\s+', ' ', code).strip()


def preprocess_text(text):
    text = text.lower()
    text = re.sub(r'[^\w\s]', ' ', text)
    if _nlp_ready:
        text = remove_stop_words(text)
        text = stem_text(text)
    return re.sub(r'\s+', ' ', text).strip()


# ===========================================================================
# 7. SPLITTING
# ===========================================================================

def split_into_paragraphs(text):
    paragraphs = re.split(r'\n\s*\n', text.strip())
    if len(paragraphs) <= 1 and len(text) > 200:
        sentences = re.split(r'[.!?]+', text)
        paragraphs, current = [], []
        for sent in sentences:
            current.append(sent.strip())
            if len(' '.join(current)) >= 100:
                paragraphs.append(' '.join(current)); current = []
        if current:
            paragraphs.append(' '.join(current))
    return [p.strip() for p in paragraphs if len(p.strip()) >= PARAGRAPH_MIN_LENGTH]


def split_code_into_blocks(code):
    blocks = []
    try:
        tree = ast.parse(code)
        ranges = [(n.lineno-1, n.end_lineno or n.lineno) for n in ast.walk(tree)
                  if isinstance(n, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef))]
        ranges.sort()
        lines = code.split('\n')
        covered = set()
        for s, e in ranges:
            b = [lines[i] for i in range(s, min(e, len(lines))) if i not in covered]
            if b:
                blocks.append('\n'.join(b)); covered.update(range(s, min(e, len(lines))))
        pre = [lines[i] for i in range(ranges[0][0] if ranges else 0) if i not in covered]
        if pre:
            blocks.insert(0, '\n'.join(pre))
    except (SyntaxError, Exception):
        blocks = [b.strip() for b in re.split(r'\n\s*\n', code) if len(b.strip()) >= 20]
    return blocks or [code]


# ===========================================================================
# 8. AST
# ===========================================================================

def code_to_ast_features(code):
    try:
        tree = ast.parse(code)
    except (SyntaxError, Exception):
        return None
    features = []
    class V(ast.NodeVisitor):
        def generic_visit(self, node):
            features.append(f"{node.__class__.__name__}({len(list(ast.iter_child_nodes(node)))})")
            if isinstance(node, ast.Call):
                if isinstance(node.func, ast.Name): features.append(f"CALL({node.func.id})")
                elif isinstance(node.func, ast.Attribute): features.append(f"CALL_ATTR({node.func.attr})")
            if isinstance(node, ast.BinOp): features.append(f"BINOP({node.op.__class__.__name__})")
            if isinstance(node, (ast.For, ast.While)): features.append(f"LOOP({node.__class__.__name__})")
            if isinstance(node, ast.If): features.append("COND(if)")
            if isinstance(node, ast.FunctionDef): features.append(f"FUNC({node.name})")
            if isinstance(node, ast.ClassDef): features.append(f"CLASS({node.name})")
            ast.NodeVisitor.generic_visit(self, node)
    V().visit(tree)
    return " ".join(features)


def code_tokens_sequence(code):
    try:
        tree = ast.parse(code)
    except (SyntaxError, Exception):
        return " ".join(re.findall(r'\b\w+\b|[+\-*/=<>!&|^~%]+|[{}()\[\];,]', code))
    tokens = []
    class V(ast.NodeVisitor):
        def generic_visit(self, node):
            if isinstance(node, ast.FunctionDef): tokens.extend(["DEF",node.name])
            elif isinstance(node, ast.ClassDef): tokens.extend(["CLASS",node.name])
            elif isinstance(node, ast.For): tokens.append("FOR")
            elif isinstance(node, ast.While): tokens.append("WHILE")
            elif isinstance(node, ast.If): tokens.append("IF")
            elif isinstance(node, ast.Return): tokens.append("RETURN")
            elif isinstance(node, ast.BinOp): tokens.append(f"OP_{node.op.__class__.__name__}")
            elif isinstance(node, ast.Compare): tokens.append("COMPARE")
            elif isinstance(node, ast.Call): tokens.append("CALL")
            elif isinstance(node, ast.Assign): tokens.append("ASSIGN")
            ast.NodeVisitor.generic_visit(self, node)
    V().visit(tree)
    return " ".join(tokens)


# ===========================================================================
# 9. WINNOWING
# ===========================================================================

def ngram_hashes(text, n=WINNOWING_WINDOW):
    if len(text) < n:
        return [hashlib.md5(text.encode()).hexdigest()] if text else []
    return [hashlib.md5(text[i:i+n].encode()).hexdigest() for i in range(len(text)-n+1)]

def winnow_select(hashes, w=WINNOWING_WINDOW):
    if not hashes or len(hashes) < w:
        return set(hashes)
    s = set()
    for i in range(len(hashes)-w+1):
        wh = hashes[i:i+w]
        s.add(hashes[i+wh.index(min(wh))])
    return s

def winnowing_similarity(t1, t2):
    s1, s2 = winnow_select(ngram_hashes(t1)), winnow_select(ngram_hashes(t2))
    if not s1 and not s2: return 1.0
    if not s1 or not s2: return 0.0
    return len(s1&s2)/len(s1|s2)


# ===========================================================================
# 10. LCS
# ===========================================================================

def lcs_ratio(t1, t2):
    if not t1 and not t2: return 1.0
    if not t1 or not t2: return 0.0
    w1, w2 = t1.split()[:2000], t2.split()[:2000]
    m, n = len(w1), len(w2)
    if m*n > 4_000_000:
        return difflib.SequenceMatcher(None, w1, w2).ratio()
    dp = [[0]*(n+1) for _ in range(m+1)]
    for i in range(1,m+1):
        for j in range(1,n+1):
            dp[i][j] = dp[i-1][j-1]+1 if w1[i-1]==w2[j-1] else max(dp[i-1][j],dp[i][j-1])
    return (2*dp[m][n])/(m+n)


# ===========================================================================
# 11. SÉMANTIQUE
# ===========================================================================

def compute_semantic_similarity(texts, new_text):
    model = get_semantic_model()
    if not model: return [0.0]*len(texts)
    emb = model.encode([new_text]+texts, convert_to_numpy=True, normalize_embeddings=True)
    return np.dot(emb[0], emb[1:].T).tolist()


# ===========================================================================
# 12. PARAGRAPH-LEVEL
# ===========================================================================

def compute_paragraph_scores(new_secs, existing_secs):
    if not new_secs or not existing_secs:
        return [0.0]*len(new_secs)
    scores = []
    for sec in new_secs:
        best = 0.0
        s1 = preprocess_text(sec)
        for es in existing_secs:
            s2 = preprocess_text(es)
            try:
                v = TfidfVectorizer(ngram_range=(1,2), min_df=1)
                mat = v.fit_transform([s1,s2])
                sim = cosine_similarity(mat[0:1],mat[1:])[0][0]
            except ValueError:
                sim = 0.0
            if sim > best: best = sim
        scores.append(round(best, 4))
    return scores


# ===========================================================================
# 13. COMBINAISON
# ===========================================================================

def combine(scores, weights):
    active, total = {}, 0.0
    for e, w in weights.items():
        if scores.get(e, -1.0) >= 0:
            active[e] = w; total += w
    if not active: return 0.0, scores
    final, detail = 0.0, {}
    for e, w in active.items():
        nw = w/total; c = scores[e]*nw; final += c
        detail[e] = {"raw": round(scores[e],4), "weight": round(nw,3), "contribution": round(c,4)}
    return round(min(final,1.0),4), detail

def classify(score):
    if score >= 0.85: return "critical"
    if score >= 0.65: return "high"
    if score >= 0.45: return "medium"
    if score >= 0.25: return "low"
    return "none"


# ===========================================================================
# 14. PIPELINE COMPLET
# ===========================================================================

def detect_text_code(processed, ast_features, token_seq, raw, sections, file_type,
                     existing, use_semantic, use_winnowing, use_ast, use_lcs):
    weights = WEIGHTS_CODE if file_type == "code" else WEIGHTS_TEXT
    results = []

    # Filtrer les soumissions de même type (text+code ensemble, pas les images)
    relevant = [s for s in existing if s["file_type"] in ("text", "code")]
    if not relevant:
        return []

    etxts = [s["processed"] for s in relevant]
    try:
        v = TfidfVectorizer(ngram_range=(1,3), min_df=1)
        m = v.fit_transform([processed]+etxts)
        tfidf_scores = cosine_similarity(m[0:1],m[1:]).flatten()
    except ValueError:
        tfidf_scores = np.zeros(len(relevant))

    sem_scores = compute_semantic_similarity(etxts, processed) if use_semantic else None
    wn_scores = [winnowing_similarity(processed, s["processed"]) for s in relevant] if use_winnowing else None

    ast_scores = []
    if use_ast and file_type=="code" and ast_features:
        for s in relevant:
            sa = s.get("ast_features","")
            if sa:
                try:
                    vv = TfidfVectorizer(ngram_range=(1,2), min_df=1)
                    mm = vv.fit_transform([ast_features,sa])
                    ast_scores.append(float(cosine_similarity(mm[0:1],mm[1:])[0][0]))
                except ValueError:
                    ast_scores.append(0.0)
            else: ast_scores.append(-1.0)
    else:
        ast_scores = [-1.0]*len(relevant)

    lcs_scores = []
    if use_lcs:
        for s in relevant:
            seq = s.get("token_sequence","")
            lcs_scores.append(lcs_ratio(token_seq, seq) if (token_seq and seq) else lcs_ratio(processed, s["processed"]))
    else:
        lcs_scores = [0.0]*len(relevant)

    for i, s in enumerate(relevant):
        sd = {
            "tfidf": float(tfidf_scores[i]),
            "semantic": float(sem_scores[i]) if sem_scores else -1.0,
            "winnowing": float(wn_scores[i]) if wn_scores else -1.0,
            "ast": float(ast_scores[i]),
            "lcs": float(lcs_scores[i]),
        }
        combined, detail = combine(sd, weights)
        p_scores = compute_paragraph_scores(sections, s.get("sections",[])) if sections else []
        suspicious = [{"section_index":idx,"score":sc,"level":classify(sc),
                       "preview":(sections[idx][:200] if idx<len(sections) else "")}
                      for idx,sc in enumerate(p_scores) if sc>=0.45]
        results.append({
            "submission_id": s["id"], "filename": s["filename"],
            "file_type": s["file_type"], "combined_score": combined,
            "level": classify(combined), "engines": detail,
            "paragraph_scores": p_scores, "suspicious_sections": suspicious,
        })

    results.sort(key=lambda x: x["combined_score"], reverse=True)
    return [r for r in results if r["combined_score"] >= 0.10][:MAX_RESULTS]


# ===========================================================================
# 15. BASE DE DONNÉES
# ===========================================================================

def load_submissions():
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE,'r',encoding='utf-8') as f:
            return json.load(f)
    return []

def save_submissions(subs):
    with open(DATA_FILE,'w',encoding='utf-8') as f:
        json.dump(subs, f, indent=2, ensure_ascii=False)


def add_submission(filename, content, file_type, metadata=None,
                   images=None, sections=None, processed=None,
                   ast_features=None, token_sequence=None):
    subs = load_submissions()
    if processed is None:
        if file_type == "code":
            processed = preprocess_code(content)
            ast_features = code_to_ast_features(content)
            token_sequence = code_tokens_sequence(content)
            sections = split_code_into_blocks(content)
        else:
            processed = preprocess_text(content)
            sections = split_into_paragraphs(content)
    else:
        sections = sections or []
        ast_features = ast_features
        token_sequence = token_sequence

    # Calculer hash + features pour chaque image
    image_data = []
    if images:
        for idx, img in enumerate(images):
            h = compute_image_hash(img)
            feat = compute_image_features(img)
            image_data.append({
                "hash": h,
                "features": feat.tolist() if feat is not None else None,
                "image_index": idx,
            })

    sub = {
        "id": hashlib.md5(f"{filename}_{len(subs)}_{content[:100]}".encode()).hexdigest()[:12],
        "filename": filename,
        "original_content": content[:2000],  # garder plus pour référence
        "processed": processed,
        "file_type": file_type,
        "ast_features": ast_features,
        "token_sequence": token_sequence,
        "sections": sections,
        "images": image_data,
        "image_count": len(image_data) if image_data else 0,
        "content_length": len(content),
        "metadata": metadata or {},
    }
    subs.append(sub)
    save_submissions(subs)
    return sub["id"]


# ===========================================================================
# 16. ENDPOINTS — JSON pour Laravel
# ===========================================================================

@app.post("/api/check")
async def api_check(
    file: Optional[UploadFile] = File(None),
    text: Optional[str] = Form(None),
    filename: Optional[str] = Form(None),
    file_type: str = Form("auto"),
    use_semantic: bool = Form(True),
    use_winnowing: bool = Form(True),
    use_ast: bool = Form(True),
    use_lcs: bool = Form(True),
):
    """
    Endpoint principal.

    Supporte : .txt, .md, .py, .js, .pdf, .docx, .png, .jpg, .jpeg, .gif, .bmp, .zip
    file_type="auto" detecte automatiquement le type.

    Retourne JSON :
    {
      "success": true,
      "data": {
        "filename": "...",
        "detected_type": "text|code|image",
        "content_analysis": { ... scores texte/code ... },
        "image_analysis": { ... scores images ... },
        "plagiarism_detected": true,
        "overall_score": 0.87,
        "overall_level": "critical"
      }
    }
    """
    if not file and not text:
        return JSONResponse(status_code=400, content={
            "success": False, "message": "Fournir 'file' ou 'text'"
        })

    # === 1. Extraire le contenu ===
    extracted_images = []
    raw_text = ""
    detected_type = file_type
    extraction_metadata = {}

    if file:
        file_bytes = await file.read()
        original_filename = file.filename or "file"
        ext = os.path.splitext(original_filename)[1].lower()

        # Si file_type="auto", on extrait automatiquement
        if file_type == "auto":
            extraction = extract_file(file_bytes, original_filename)
            raw_text = extraction["text"]
            extracted_images = extraction.get("images", [])
            detected_type = extraction["file_type"]
            extraction_metadata = extraction.get("metadata", {})
        else:
            # Force le type spécifié
            try:
                raw_text = file_bytes.decode('utf-8', errors='ignore')
            except Exception:
                raw_text = file_bytes.decode('latin-1', errors='ignore')
    else:
        raw_text = text
        original_filename = filename or "inline.txt"
        detected_type = file_type if file_type != "auto" else "text"

    # Si c'est une image directe sans texte, on ne fait que la comparaison d'images
    if detected_type == "image" and not raw_text:
        existing = load_submissions()

        # Collecter toutes les images de la base
        all_existing_images = []
        for s in existing:
            for img_data in s.get("images", []):
                entry = {**img_data, "filename": s["filename"]}
                all_existing_images.append(entry)

        # Comparer les images
        image_matches = compare_images(extracted_images, all_existing_images)

        max_img_score = max([m["confidence"] for m in image_matches], default=0)
        max_img_level = image_matches[0]["level"] if image_matches else "none"

        return JSONResponse(content={
            "success": True,
            "data": {
                "filename": original_filename,
                "detected_type": "image",
                "extraction": extraction_metadata,
                "content_analysis": {
                    "analyzed": False,
                    "reason": "Image seule sans texte",
                    "text_matches": [],
                },
                "image_analysis": {
                    "analyzed": True,
                    "images_checked": len(extracted_images),
                    "images_in_database": len(all_existing_images),
                    "image_matches": image_matches,
                    "max_score": max_img_score,
                    "max_level": max_img_level,
                },
                "plagiarism_detected": max_img_score >= 0.70,
                "overall_score": max_img_score,
                "overall_level": max_img_level,
            }
        })

    # === 2. Analyser le texte/code ===
    if detected_type == "code":
        processed = preprocess_code(raw_text)
        ast_f = code_to_ast_features(raw_text)
        tok_seq = code_tokens_sequence(raw_text)
        sections = split_code_into_blocks(raw_text)
    else:
        processed = preprocess_text(raw_text)
        ast_f, tok_seq = None, None
        sections = split_into_paragraphs(raw_text)

    existing = load_submissions()

    text_matches = []
    if processed and existing:
        text_matches = detect_text_code(
            processed=processed, ast_features=ast_f, token_seq=tok_seq,
            raw=raw_text, sections=sections, file_type=detected_type,
            existing=existing, use_semantic=use_semantic and detected_type=="text",
            use_winnowing=use_winnowing, use_ast=use_ast and detected_type=="code",
            use_lcs=use_lcs,
        )

    max_text_score = text_matches[0]["combined_score"] if text_matches else 0
    max_text_level = text_matches[0]["level"] if text_matches else "none"

    # === 3. Analyser les images ===
    image_matches = []
    max_img_score = 0
    max_img_level = "none"

    if extracted_images and existing:
        all_existing_images = []
        for s in existing:
            for img_data in s.get("images", []):
                entry = {**img_data, "filename": s["filename"]}
                all_existing_images.append(entry)

        image_matches = compare_images(extracted_images, all_existing_images)
        max_img_score = max([m["confidence"] for m in image_matches], default=0)
        max_img_level = image_matches[0]["level"] if image_matches else "none"

    # === 4. Score global combiné ===
    if max_img_score > 0 and max_text_score > 0:
        overall_score = round(max(max_text_score, max_img_score), 4)
        levels_order = {"critical":4,"high":3,"medium":2,"low":1,"none":0}
        overall_level = max([max_text_level, max_img_level], key=lambda x: levels_order.get(x,0))
    elif max_img_score > 0:
        overall_score = max_img_score
        overall_level = max_img_level
    else:
        overall_score = max_text_score
        overall_level = max_text_level

    # === 5. Réponse JSON ===
    levels = Counter(m["level"] for m in text_matches)
    img_levels = Counter(m["level"] for m in image_matches)

    return JSONResponse(content={
        "success": True,
        "data": {
            "filename": original_filename,
            "detected_type": detected_type,
            "extraction": extraction_metadata,
            "content_length": len(raw_text),
            "images_extracted": len(extracted_images),
            "num_comparisons": len(existing),

            "content_analysis": {
                "analyzed": bool(processed and existing),
                "max_score": max_text_score,
                "max_level": max_text_level,
                "summary": {
                    "critical": levels.get("critical",0),
                    "high": levels.get("high",0),
                    "medium": levels.get("medium",0),
                    "low": levels.get("low",0),
                },
                "paragraphs_analysed": len(sections),
                "text_matches": text_matches,
            },

            "image_analysis": {
                "analyzed": bool(extracted_images and existing),
                "images_checked": len(extracted_images),
                "images_in_database": sum(len(s.get("images",[])) for s in existing),
                "summary": {
                    "critical": img_levels.get("critical",0),
                    "high": img_levels.get("high",0),
                    "medium": img_levels.get("medium",0),
                    "low": img_levels.get("low",0),
                },
                "image_matches": image_matches,
                "max_score": max_img_score,
                "max_level": max_img_level,
            },

            "plagiarism_detected": overall_score >= 0.25,
            "overall_score": overall_score,
            "overall_level": overall_level,
        }
    })


@app.post("/api/add")
async def api_add(file: UploadFile = File(...), file_type: str = Form("auto"),
                  metadata: Optional[str] = Form(None)):
    """Ajoute un fichier (texte, code, PDF, Word, image) à la base."""
    file_bytes = await file.read()
    filename = file.filename

    # Extraction automatique
    extraction = extract_file(file_bytes, filename)
    content = extraction["text"]
    images = extraction.get("images", [])
    ftype = extraction["file_type"]
    meta = extraction.get("metadata", {})
    if metadata:
        try: meta.update(json.loads(metadata))
        except: pass

    # Ajouter à la base
    sub_id = add_submission(
        filename=filename, content=content, file_type=ftype,
        metadata=meta, images=images,
    )

    return JSONResponse(content={
        "success": True,
        "id": sub_id,
        "filename": filename,
        "type": ftype,
        "images_stored": len(images),
        "content_length": len(content),
    })


@app.post("/api/add/text")
async def api_add_text(text: str = Form(...), filename: str = Form("text.txt"),
                        metadata: Optional[str] = Form(None)):
    meta = json.loads(metadata) if metadata else {}
    sub_id = add_submission(filename, text, "text", meta)
    return JSONResponse(content={"success":True, "id":sub_id, "filename":filename})


@app.post("/api/add/code")
async def api_add_code(code: str = Form(...), filename: str = Form("code.py"),
                        metadata: Optional[str] = Form(None)):
    meta = json.loads(metadata) if metadata else {}
    sub_id = add_submission(filename, code, "code", meta)
    return JSONResponse(content={"success":True, "id":sub_id, "filename":filename})


@app.get("/api/stats")
async def api_stats():
    subs = load_submissions()
    return JSONResponse(content={
        "success": True,
        "data": {
            "total": len(subs),
            "text": sum(1 for s in subs if s["file_type"]=="text"),
            "code": sum(1 for s in subs if s["file_type"]=="code"),
            "image": sum(1 for s in subs if s["file_type"]=="image"),
            "total_images": sum(s.get("image_count",0) for s in subs),
            "total_chars": sum(s.get("content_length",0) for s in subs),
        }
    })


@app.get("/api/list")
async def api_list(limit: int = 50, file_type: Optional[str] = None):
    subs = load_submissions()
    if file_type:
        subs = [s for s in subs if s["file_type"]==file_type]
    return JSONResponse(content={
        "success": True,
        "data": [{"id":s["id"],"filename":s["filename"],"file_type":s["file_type"],
                   "length":s.get("content_length",0),"images":s.get("image_count",0)}
                  for s in subs[:limit]]
    })


@app.delete("/api/delete/{sid}")
async def api_delete(sid: str):
    subs = load_submissions()
    n = len(subs)
    subs = [s for s in subs if s["id"]!=sid]
    if len(subs)==n:
        return JSONResponse(status_code=404, content={"success":False,"message":"Non trouve"})
    save_submissions(subs)
    return JSONResponse(content={"success":True,"id":sid})


@app.get("/api/health")
async def api_health():
    return JSONResponse(content={
        "success": True,
        "data": {
            "status": "healthy",
            "version": "4.0.0",
            "nlp": _nlp_ready,
            "semantic": get_semantic_model() is not None,
            "pdf_extraction": HAS_FITZ,
            "word_extraction": HAS_DOCX,
            "image_comparison": HAS_PIL and HAS_IMAGEHASH,
            "corpus_size": len(load_submissions()),
        }
    })


# ===========================================================================
# 17. ZIP AVANCÉ — Analyse fichier par fichier
# ===========================================================================

# Extensions reconnues pour l'extraction individuelle
CODE_EXTS_SET = {
    '.py', '.js', '.jsx', '.ts', '.tsx', '.java', '.c', '.cpp', '.h', '.hpp',
    '.cs', '.php', '.rb', '.go', '.rs', '.swift', '.kt', '.scala', '.r',
    '.sql', '.sh', '.bash', '.html', '.css', '.scss', '.xml', '.yaml', '.yml',
    '.json', '.toml', '.ini', '.cfg', '.conf',
}
TEXT_EXTS_SET = {'.txt', '.md', '.rst', '.csv'}
IMAGE_EXTS_SET = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.tiff', '.tif'}
DOC_EXTS_SET = {'.pdf', '.docx'}

# Seuil minimum de contenu pour analyser un fichier
MIN_CONTENT_LENGTH = 50


def extract_zip_individual(zip_bytes: bytes) -> List[Dict]:
    """
    Extrait chaque fichier du ZIP individuellement.
    Retourne une liste de dicts, un par fichier extrait :
      [{
          "filename": "src/main.py",
          "basename": "main.py",
          "extension": ".py",
          "file_type": "code",
          "text": "...contenu...",
          "images": [PIL.Image, ...],
          "content_length": 1234,
      }, ...]
    """
    files = []
    try:
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
            for name in zf.namelist():
                # Ignorer dossiers, fichiers cachés macOS
                if name.startswith('__MACOSX') or name.endswith('/'):
                    continue
                # Ignorer fichiers trop petits ou invisibles
                basename = os.path.basename(name)
                if basename.startswith('.') and basename != '.':
                    continue

                try:
                    data = zf.read(name)
                    ext = os.path.splitext(name)[1].lower()

                    entry = {
                        "filename": name,
                        "basename": basename,
                        "extension": ext,
                        "text": "",
                        "images": [],
                        "file_type": "text",
                        "content_length": 0,
                    }

                    # === Image ===
                    if ext in IMAGE_EXTS_SET and HAS_PIL:
                        try:
                            img = Image.open(io.BytesIO(data))
                            if img.mode != 'RGB':
                                img = img.convert('RGB')
                            entry["images"] = [img]
                            entry["file_type"] = "image"
                            entry["content_length"] = len(data)
                        except Exception:
                            continue

                    # === PDF ===
                    elif ext == '.pdf' and HAS_FITZ:
                        tmp = None
                        try:
                            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as t:
                                t.write(data); tmp = t.name
                            result = extract_pdf(tmp)
                            entry["text"] = result.get("text", "")
                            entry["images"] = result.get("images", [])
                            entry["file_type"] = "text"
                            entry["content_length"] = len(entry["text"])
                        except Exception:
                            continue
                        finally:
                            if tmp and os.path.exists(tmp): os.unlink(tmp)

                    # === Word ===
                    elif ext == '.docx' and HAS_DOCX:
                        tmp = None
                        try:
                            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as t:
                                t.write(data); tmp = t.name
                            result = extract_docx(tmp)
                            entry["text"] = result.get("text", "")
                            entry["images"] = result.get("images", [])
                            entry["file_type"] = "text"
                            entry["content_length"] = len(entry["text"])
                        except Exception:
                            continue
                        finally:
                            if tmp and os.path.exists(tmp): os.unlink(tmp)

                    # === Code / Texte brut ===
                    else:
                        try:
                            txt = data.decode('utf-8', errors='ignore')
                        except Exception:
                            continue
                        entry["text"] = txt
                        entry["file_type"] = "code" if ext in CODE_EXTS_SET else "text"
                        entry["content_length"] = len(txt)

                    # Garder uniquement les fichiers avec assez de contenu
                    if entry["text"] or entry["images"]:
                        files.append(entry)

                except Exception:
                    continue

    except Exception as e:
        pass

    return files


def compare_two_files(file_a: Dict, file_b: Dict,
                      use_semantic=True, use_winnowing=True, use_ast=True, use_lcs=True) -> Dict:
    """
    Compare deux fichiers entre eux et retourne un score détaillé.

    file_a, file_b : dict avec {text, file_type, filename, ...}

    Retourne :
      {
          "file_a": "main.py",
          "file_b": "utils.py",
          "combined_score": 0.87,
          "level": "high",
          "engines": {...},
      }
    """
    text_a = file_a.get("text", "")
    text_b = file_b.get("text", "")

    if not text_a.strip() or not text_b.strip():
        return {"file_a": file_a["filename"], "file_b": file_b["filename"],
                "combined_score": 0.0, "level": "none", "engines": {}}

    # Ne pas comparer un code avec du texte (pas pertinent)
    if file_a["file_type"] != file_b["file_type"]:
        return {"file_a": file_a["filename"], "file_b": file_b["filename"],
                "combined_score": 0.0, "level": "none", "engines": {}}

    is_code = file_a["file_type"] == "code"
    weights = WEIGHTS_CODE if is_code else WEIGHTS_TEXT

    # Prétraiter
    if is_code:
        proc_a = preprocess_code(text_a)
        proc_b = preprocess_code(text_b)
    else:
        proc_a = preprocess_text(text_a)
        proc_b = preprocess_text(text_b)

    if not proc_a.strip() or not proc_b.strip():
        return {"file_a": file_a["filename"], "file_b": file_b["filename"],
                "combined_score": 0.0, "level": "none", "engines": {}}

    # TF-IDF
    try:
        v = TfidfVectorizer(ngram_range=(1,3), min_df=1)
        m = v.fit_transform([proc_a, proc_b])
        tfidf_sim = float(cosine_similarity(m[0:1], m[1:])[0][0])
    except ValueError:
        tfidf_sim = 0.0

    # Sémantique
    if use_semantic and not is_code:
        model = get_semantic_model()
        if model:
            try:
                emb = model.encode([proc_a, proc_b], convert_to_numpy=True, normalize_embeddings=True)
                sem_sim = float(np.dot(emb[0], emb[1]))
            except Exception:
                sem_sim = 0.0
        else:
            sem_sim = 0.0
    elif is_code:
        # Sémantique allégée pour code
        model = get_semantic_model()
        if model:
            try:
                emb = model.encode([proc_a, proc_b], convert_to_numpy=True, normalize_embeddings=True)
                sem_sim = float(np.dot(emb[0], emb[1]))
            except Exception:
                sem_sim = 0.0
        else:
            sem_sim = 0.0
    else:
        sem_sim = -1.0

    # Winnowing
    wn_sim = winnowing_similarity(proc_a, proc_b) if use_winnowing else -1.0

    # AST (code uniquement)
    ast_sim = -1.0
    if use_ast and is_code:
        ast_a = code_to_ast_features(text_a)
        ast_b = code_to_ast_features(text_b)
        if ast_a and ast_b:
            try:
                vv = TfidfVectorizer(ngram_range=(1,2), min_df=1)
                mm = vv.fit_transform([ast_a, ast_b])
                ast_sim = float(cosine_similarity(mm[0:1], mm[1:])[0][0])
            except ValueError:
                ast_sim = 0.0

    # LCS
    if use_lcs:
        if is_code:
            tok_a = code_tokens_sequence(text_a)
            tok_b = code_tokens_sequence(text_b)
            lcs_sim = lcs_ratio(tok_a, tok_b) if (tok_a and tok_b) else lcs_ratio(proc_a, proc_b)
        else:
            lcs_sim = lcs_ratio(proc_a, proc_b)
    else:
        lcs_sim = 0.0

    # Combiner
    scores = {
        "tfidf": tfidf_sim,
        "semantic": sem_sim,
        "winnowing": wn_sim,
        "ast": ast_sim,
        "lcs": lcs_sim,
    }
    combined, detail = combine(scores, weights)

    return {
        "file_a": file_a["filename"],
        "file_b": file_b["filename"],
        "combined_score": combined,
        "level": classify(combined),
        "engines": detail,
    }


@app.post("/api/check-zip")
async def api_check_zip(
    file: UploadFile = File(...),
    use_semantic: bool = Form(True),
    use_winnowing: bool = Form(True),
    use_ast: bool = Form(True),
    use_lcs: bool = Form(True),
    cross_compare: bool = Form(True),
    add_to_database: bool = Form(True),
):
    """
    Endpoint avancé pour analyser un ZIP.
    Contrairement à /api/check qui fusionne tout le contenu,
    cet endpoint analyse CHAQUE fichier individuellement, puis
    compare les fichiers ENTRE EUX dans le ZIP.

    Retourne JSON :
    {
      "success": true,
      "data": {
        "filename": "projet.zip",
        "files_extracted": [
            {"filename": "main.py", "file_type": "code", "content_length": 1234, ...},
            ...
        ],
        "cross_file_analysis": [
            {
                "file_a": "main.py",
                "file_b": "utils.py",
                "combined_score": 0.87,
                "level": "high",
                "engines": {...}
            },
            ...
        ],
        "per_file_database_analysis": [
            {
                "filename": "main.py",
                "max_score": 0.65,
                "max_level": "high",
                "best_match": "existing_file.py",
                "matches": [...]
            },
            ...
        ],
        "image_analysis": {...},
        "overall_score": 0.87,
        "overall_level": "critical",
        "plagiarism_detected": true
      }
    }
    """
    init_nlp()
    get_semantic_model()

    zip_bytes = await file.read()
    original_filename = file.filename or "archive.zip"

    # Vérifier que c'est bien un ZIP
    try:
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
            pass
    except zipfile.BadZipFile:
        return JSONResponse(status_code=400, content={
            "success": False,
            "message": "Le fichier n'est pas un ZIP valide."
        })

    # === 1. Extraire chaque fichier individuellement ===
    extracted_files = extract_zip_individual(zip_bytes)

    if not extracted_files:
        return JSONResponse(content={
            "success": True,
            "data": {
                "filename": original_filename,
                "files_extracted": [],
                "cross_file_analysis": [],
                "per_file_database_analysis": [],
                "image_analysis": {"analyzed": False, "image_matches": []},
                "overall_score": 0.0,
                "overall_level": "none",
                "plagiarism_detected": False,
                "message": "Aucun fichier analysable trouvé dans le ZIP.",
            }
        })

    # Filtrer les fichiers texte/code (pas les images pures pour le texte)
    text_code_files = [f for f in extracted_files if f["file_type"] in ("text", "code") and len(f.get("text", "")) >= MIN_CONTENT_LENGTH]
    image_files = [f for f in extracted_files if f["images"]]

    # === 2. Comparaison CROISÉE entre fichiers du ZIP ===
    cross_matches = []
    if cross_compare and len(text_code_files) > 1:
        for i in range(len(text_code_files)):
            for j in range(i + 1, len(text_code_files)):
                result = compare_two_files(
                    text_code_files[i], text_code_files[j],
                    use_semantic=use_semantic,
                    use_winnowing=use_winnowing,
                    use_ast=use_ast,
                    use_lcs=use_lcs,
                )
                if result["combined_score"] >= 0.10:
                    cross_matches.append(result)

        cross_matches.sort(key=lambda x: x["combined_score"], reverse=True)
        # Limiter à 20 résultats croisés
        cross_matches = cross_matches[:20]

    # === 3. Comparaison de chaque fichier avec la BASE DE DONNÉES ===
    existing = load_submissions()
    per_file_results = []

    for f in text_code_files:
        # Prétraiter
        if f["file_type"] == "code":
            processed = preprocess_code(f["text"])
            ast_f = code_to_ast_features(f["text"])
            tok_seq = code_tokens_sequence(f["text"])
            sections = split_code_into_blocks(f["text"])
        else:
            processed = preprocess_text(f["text"])
            ast_f = None
            tok_seq = None
            sections = split_into_paragraphs(f["text"])

        if not processed.strip():
            per_file_results.append({
                "filename": f["filename"],
                "file_type": f["file_type"],
                "max_score": 0.0,
                "max_level": "none",
                "best_match": None,
                "matches": [],
                "skipped": True,
                "reason": "Contenu trop court ou vide après prétraitement",
            })
            continue

        # Comparer avec la base
        matches = detect_text_code(
            processed=processed,
            ast_features=ast_f,
            token_seq=tok_seq,
            raw=f["text"],
            sections=sections,
            file_type=f["file_type"],
            existing=existing,
            use_semantic=use_semantic and f["file_type"] == "text",
            use_winnowing=use_winnowing,
            use_ast=use_ast and f["file_type"] == "code",
            use_lcs=use_lcs,
        )

        max_score = matches[0]["combined_score"] if matches else 0
        max_level = matches[0]["level"] if matches else "none"
        best_match = matches[0]["filename"] if matches else None

        per_file_results.append({
            "filename": f["filename"],
            "file_type": f["file_type"],
            "content_length": f["content_length"],
            "max_score": max_score,
            "max_level": max_level,
            "best_match": best_match,
            "matches": matches,
        })

        # Ajouter à la base de données (optionnel)
        if add_to_database:
            try:
                add_submission(
                    filename=f"zip://{original_filename}/{f['filename']}",
                    content=f["text"],
                    file_type=f["file_type"],
                    metadata={
                        "source_zip": original_filename,
                        "zip_path": f["filename"],
                    },
                    images=f.get("images", []),
                )
            except Exception:
                pass

    # === 4. Comparaison d'images (interne + base) ===
    all_zip_images = []
    for f in image_files:
        for idx, img in enumerate(f["images"]):
            all_zip_images.append(img)

    image_matches = []
    img_max_score = 0.0
    img_max_level = "none"

    if all_zip_images and existing:
        # Comparer images du ZIP avec la base
        all_existing_images = []
        for s in existing:
            for img_data in s.get("images", []):
                entry = {**img_data, "filename": s["filename"]}
                all_existing_images.append(entry)

        image_matches = compare_images(all_zip_images, all_existing_images)
        img_max_score = max([m["confidence"] for m in image_matches], default=0)
        img_max_level = image_matches[0]["level"] if image_matches else "none"

    # Comparer images entre elles dans le ZIP
    zip_image_matches = []
    if len(all_zip_images) > 1 and HAS_IMAGEHASH:
        for i in range(len(all_zip_images)):
            for j in range(i + 1, len(all_zip_images)):
                h_i = compute_image_hash(all_zip_images[i])
                h_j = compute_image_hash(all_zip_images[j])
                if h_i and h_j and HAS_IMAGEHASH:
                    try:
                        hi = imagehash.hex_to_hash(h_i, hash_size=16)
                        hj = imagehash.hex_to_hash(h_j, hash_size=16)
                        dist = hi - hj
                        if dist <= IMAGE_HASH_THRESHOLD:
                            conf = 1.0 - (dist / 64.0)
                            level = "critical" if dist <= 5 else ("high" if dist <= 10 else "medium")
                            zip_image_matches.append({
                                "image_a_index": i,
                                "image_b_index": j,
                                "phash_distance": dist,
                                "confidence": round(conf, 4),
                                "level": level,
                            })
                    except Exception:
                        pass

        # Compléter image matches avec les correspondances internes au ZIP
        for zm in zip_image_matches:
            image_matches.append({
                "new_image_index": zm["image_a_index"],
                "matched_filename": f"[ZIP interne] image #{zm['image_b_index']+1}",
                "matched_image_index": zm["image_b_index"],
                "level": zm["level"],
                "confidence": zm["confidence"],
                "phash_distance": zm["phash_distance"],
                "feature_similarity": None,
                "source": "cross_zip",
            })
            img_max_score = max(img_max_score, zm["confidence"])
            levels_order = {"critical": 4, "high": 3, "medium": 2, "low": 1, "none": 0}
            if levels_order.get(zm["level"], 0) > levels_order.get(img_max_level, 0):
                img_max_level = zm["level"]

    # === 5. Score global ===
    # Prendre le max entre : cross-file, per-file max, images
    cross_max = cross_matches[0]["combined_score"] if cross_matches else 0
    per_file_max = max([r["max_score"] for r in per_file_results], default=0)

    overall_score = round(max(cross_max, per_file_max, img_max_score), 4)

    levels_order = {"critical": 4, "high": 3, "medium": 2, "low": 1, "none": 0}
    cross_level = cross_matches[0]["level"] if cross_matches else "none"
    per_level = max([r["max_level"] for r in per_file_results], key=lambda x: levels_order.get(x, 0), default="none") if per_file_results else "none"
    overall_level = max([cross_level, per_level, img_max_level], key=lambda x: levels_order.get(x, 0))

    # Résumé des niveaux
    all_matches_combined = cross_matches + [m for r in per_file_results for m in r.get("matches", [])]
    summary = Counter(m["level"] for m in all_matches_combined)

    # Résumé par fichier
    file_summary = []
    for r in per_file_results:
        file_summary.append({
            "filename": r["filename"],
            "file_type": r["file_type"],
            "max_score": r["max_score"],
            "level": r["max_level"],
            "best_match": r["best_match"],
        })

    return JSONResponse(content={
        "success": True,
        "data": {
            "filename": original_filename,

            "zip_info": {
                "total_files": len(extracted_files),
                "text_code_files": len(text_code_files),
                "image_files": len(image_files),
                "files_skipped": len(extracted_files) - len(text_code_files) - len(image_files),
                "files_list": [f["filename"] for f in extracted_files],
            },

            "files_extracted": [
                {
                    "filename": f["filename"],
                    "file_type": f["file_type"],
                    "extension": f["extension"],
                    "content_length": f["content_length"],
                    "images_count": len(f.get("images", [])),
                }
                for f in extracted_files
            ],

            # Comparaison croisée entre fichiers du ZIP
            "cross_file_analysis": {
                "analyzed": len(text_code_files) > 1 and cross_compare,
                "pairs_compared": len(text_code_files) * (len(text_code_files) - 1) // 2 if cross_compare else 0,
                "matches_found": len(cross_matches),
                "max_score": cross_max,
                "max_level": cross_level,
                "summary": {
                    "critical": summary.get("critical", 0),
                    "high": summary.get("high", 0),
                    "medium": summary.get("medium", 0),
                    "low": summary.get("low", 0),
                },
                "matches": cross_matches,
            },

            # Comparaison de chaque fichier avec la base
            "per_file_database_analysis": {
                "files_analyzed": len([r for r in per_file_results if not r.get("skipped")]),
                "files_skipped": len([r for r in per_file_results if r.get("skipped")]),
                "max_score": per_file_max,
                "max_level": per_level,
                "results": per_file_results,
            },

            # Comparaison d'images
            "image_analysis": {
                "analyzed": bool(all_zip_images and existing),
                "images_in_zip": len(all_zip_images),
                "images_in_database": sum(len(s.get("images", [])) for s in existing),
                "matches_found": len(image_matches),
                "max_score": img_max_score,
                "max_level": img_max_level,
                "image_matches": image_matches,
            },

            # Score global
            "plagiarism_detected": overall_score >= 0.25,
            "overall_score": overall_score,
            "overall_level": overall_level,
        }
    })


# ===========================================================================
# 18. LANCEMENT
# ===========================================================================

if __name__ == "__main__":
    init_nlp()
    get_semantic_model()

    print("=" * 60)
    print("  Plagiarism API v4 — Full Multi-Format + Images")
    print("=" * 60)
    print()
    print("Formats supportes :")
    print("  Texte  : .txt .md")
    print("  Code   : .py .js .java .c .cpp .php ...")
    print("  PDF    : .pdf (texte + images)")
    print("  Word   : .docx (texte + images)")
    print("  Images : .png .jpg .jpeg .gif .bmp .webp")
    print("  Archive: .zip (multi-fichiers)")
    print()
    print("Detection :")
    print("  Texte : TF-IDF + BERT + Winnowing + LCS + NLTK FR")
    print("  Code  : TF-IDF + AST + Winnowing + LCS")
    print("  Images: pHash + Feature comparison")
    print()
    print("Endpoints :")
    print("  POST /api/check       — Analyse complete (JSON)")
    print("  POST /api/add         — Ajouter fichier (auto-detect)")
    print("  POST /api/add/text    — Ajouter texte")
    print("  POST /api/add/code    — Ajouter code")
    print("  GET  /api/stats       — Statistiques")
    print("  GET  /api/list        — Lister documents")
    print("  GET  /api/health      — Health check")
    print()

    uvicorn.run(app, host="0.0.0.0", port=5000)
